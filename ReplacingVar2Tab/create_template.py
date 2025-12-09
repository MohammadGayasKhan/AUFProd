from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('Employee Information Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('Personal Details', level=1)
doc.add_paragraph('Employee Name: {{ EmpName }}')
doc.add_paragraph('Employee ID: {{ EmpID }}')
doc.add_paragraph('Date of Birth: {{ DOB }}')

doc.add_paragraph()

doc.add_heading('Employment Details', level=1)
doc.add_paragraph('Department: {{ Department }}')
doc.add_paragraph('Position: {{ Position }}')
doc.add_paragraph('Salary: ${{ Salary }}')
doc.add_paragraph('Join Date: {{ JoinDate }}')
doc.add_paragraph('Manager: {{ Manager }}')

doc.add_paragraph()

doc.add_heading('Educational Qualifications', level=1)
doc.add_paragraph('The qualifications are:')

doc.add_paragraph('{%- for qual in Qualifications %}\n{{ qual.year }} – {{ qual.degree }}\n{% endfor %}')

doc.add_paragraph('End of Report')

doc.save('Template.docx')
print("✓ Template created: Template.docx")
print("\nTemplate uses Jinja2 loop syntax for qualifications:")
print("  {%- for qual in Qualifications %}")
print("  {{ qual.year }} – {{ qual.degree }}")
print("  {% endfor %}")
print("\nClose Template.docx if open, then run: python change.py")




