from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading('Employee Information', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add employee basic info
doc.add_paragraph('Employee Name: {{ EmpName }}')
doc.add_paragraph('Employee ID: {{ EmpID }}')
doc.add_paragraph('Date of Birth: {{ DOB }}')

doc.add_paragraph()

# Add a heading for qualifications
doc.add_heading('Qualifications', level=1)

# Create a table with header row - this will be populated dynamically
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'

# Add header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Degree'
header_cells[1].text = 'Year'

# Make header bold
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

# Add Jinja2 loop template for dynamic rows
# This is the key part - docxtpl will duplicate this row for each item in Qualifications
row_cells = table.add_row().cells
row_cells[0].text = '{% for qual in Qualifications %}{{ qual.degree }}'
row_cells[1].text = '{{ qual.year }}{% endfor %}'

doc.add_paragraph()

# Add employment details section
doc.add_heading('Employment Details', level=1)
table2 = doc.add_table(rows=4, cols=2)
table2.style = 'Light Grid Accent 1'

employment_data = [
    ('Department', '{{ Department }}'),
    ('Position', '{{ Position }}'),
    ('Salary', '{{ Salary }}'),
    ('Manager', '{{ Manager }}')
]

for i, (label, value) in enumerate(employment_data):
    row = table2.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value

# Save the document
doc.save('Template.docx')
print("Template created: Template.docx")
print("\nTemplate includes:")
print("- Basic employee info: {{ EmpName }}, {{ EmpID }}, {{ DOB }}")
print("- Dynamic table using: {% for qual in Qualifications %}")
print("- Static employment details table")
print("\nNOTE: For proper table loops, manually edit the template:")
print("1. Open Template.docx in Word")
print("2. In the Qualifications table, select the second row")
print("3. Put {% for qual in Qualifications %} BEFORE the row")
print("4. Put {% endfor %} AFTER the row")
print("5. In the row cells, use {{ qual.degree }} and {{ qual.year }}")
