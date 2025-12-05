import json
import re
from docx import Document

def process_file(input_filename, output_filename=None):
    with open('details.json', 'r') as f:
        details = json.load(f)
    
    doc = Document(input_filename)
    
    pattern = r'\$([A-Za-z_][A-Za-z0-9_]*)'
    variables = []
    
    for paragraph in doc.paragraphs:
        found_vars = re.findall(pattern, paragraph.text)
        variables.extend(found_vars)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                found_vars = re.findall(pattern, cell.text)
                variables.extend(found_vars)
    
    unique_variables = list(dict.fromkeys(variables))
    print(f"Found variables: {unique_variables}")
    
    for paragraph in doc.paragraphs:
        for var in unique_variables:
            if var in details and f'${var}' in paragraph.text:
                for run in paragraph.runs:
                    if f'${var}' in run.text:
                        run.text = run.text.replace(f'${var}', str(details[var]))
                        print(f"Replaced ${var} with {details[var]} in paragraph")
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for var in unique_variables:
                        if var in details and f'${var}' in paragraph.text:
                            for run in paragraph.runs:
                                if f'${var}' in run.text:
                                    run.text = run.text.replace(f'${var}', str(details[var]))
                                    print(f"Replaced ${var} with {details[var]} in table")
    
    if output_filename is None:
        output_filename = input_filename.replace('.docx', '_output.docx')
    
    doc.save(output_filename)
    
    print(f"\nProcessed file saved to: {output_filename}")
    return unique_variables

if __name__ == "__main__":
    variables = process_file('template.docx', 'output.docx')
    
    print("\n" + "="*50)
    print("Extracted place holders:", variables)
    print("="*50)
