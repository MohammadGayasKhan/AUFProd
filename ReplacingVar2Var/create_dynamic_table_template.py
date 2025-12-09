"""
Template Creation Guide for Dynamic Tables with docxtpl

To create a template with dynamic table rows in Microsoft Word:

1. Create your table with headers in the first row
2. Add a data row below the headers
3. In Word, place your cursor BEFORE the data row
4. Insert a special paragraph with: {%tr for qual in Qualifications %}
5. In the data row cells, use: {{ qual.degree }} and {{ qual.year }}
6. After the data row, insert: {%tr endfor %}

Example Table Structure in Word:
┌──────────────┬──────────┐
│   Degree     │   Year   │  ← Header row
├──────────────┼──────────┤
│{%tr for qual in Qualifications %}  ← Put this BEFORE the row
├──────────────┼──────────┤
│{{ qual.degree }}│{{ qual.year }}│ ← Data row with variables
├──────────────┼──────────┤
│{%tr endfor %}               ← Put this AFTER the row
└──────────────┴──────────┘

The {%tr %} tag tells docxtpl to repeat the table row.

For your JSON structure:
{
    "EmpName": "John Doe",
    "Qualifications": [
        { "degree": "B.Tech CSE", "year": 2020 },
        { "degree": "M.Tech AIML", "year": 2023 }
    ]
}

This will generate 2 rows in the table automatically!
"""

print(__doc__)

# Create a simple template manually
from docx import Document

doc = Document()
doc.add_heading('Employee Information', 0)
doc.add_paragraph('Name: {{ EmpName }}')
doc.add_paragraph()
doc.add_heading('Qualifications', 1)

# Create table
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'

# Headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Degree'
hdr_cells[1].text = 'Year'

# Add a row with instructions
row = table.add_row().cells
row[0].text = '{%tr for qual in Qualifications %}'
row[1].text = ''

row = table.add_row().cells
row[0].text = '{{ qual.degree }}'
row[1].text = '{{ qual.year }}'

row = table.add_row().cells
row[0].text = '{%tr endfor %}'
row[1].text = ''

doc.save('Template_with_loop.docx')
print("\nCreated: Template_with_loop.docx")
print("This template has the proper {%tr for/endfor %} syntax for dynamic tables!")
